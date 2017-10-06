import os
import base64
import time

from docx import Document
from docx.shared import Inches
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

from flask_mail import Message
from flask import render_template
from extensions import mail
from flask import current_app
from flask import request


def save_signature(base64_str, name_3_emp_name, frm_name):
    path = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'static', 'uploads', name_3_emp_name)
    file_name = '{}_{}_{}.png'.format(path, frm_name, time.time())
    image = base64.b64decode(base64_str.split(',')[1])
    with open(file_name, 'wb') as f:
        f.write(image)
        f.close()
    return file_name

def send_link_as_mail(**kwargs):
    subject = 'Probation Form - {}'.format(kwargs['emp_name'])

    msg = Message(subject, sender=kwargs['reviewer_email'], recipients=[
        kwargs['emp_email']
    ])

    msg.html = """Please click on the link below to sign your probation status form.<br>
    <a href="http://{0}/document/{1}/{2}">Click here</a>
    """.format(request.host,kwargs['rev_emp_code'],kwargs['emp_code'])

    mail.send(msg)





def save_document_as_docx(**kwargs):
    document = Document()
    style = document.styles['Normal']
    font = style.font
    font.size = Pt(10)
    document.add_heading('PROBATION EVALUATION FORM')
    document.add_paragraph('Employee Information')

    paragraph = document.add_paragraph("""Employee Name : %s Employee Code : %s Date :	%s.""" % (kwargs['emp_name'],kwargs['emp_code'], kwargs['date']))
    paragraph_format = paragraph.paragraph_format
    paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph_format.line_spacing = Pt(12)

    paragraph = document.add_paragraph("""Department: %s Period of Review :	%s """ % (kwargs['department'], kwargs['period_of_review']))
    paragraph_format = paragraph.paragraph_format
    paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph_format.line_spacing = Pt(12)

    paragraph = document.add_paragraph("""Reviewer:	%s	Reviewers Title: %s """ %(kwargs['reviewer'],kwargs['reviewers_title']))
    paragraph_format = paragraph.paragraph_format
    paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph_format.line_spacing = Pt(12)

    table = document.add_table(rows=1, cols=6)

    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Performance Evaluation'
    hdr_cells[1].text = 'Excellent'
    hdr_cells[2].text = 'Good'
    hdr_cells[3].text = 'Fair'
    hdr_cells[4].text = 'Poor'
    hdr_cells[5].text = 'Comments'

    row_cells = table.add_row().cells
    row_cells[0].text = 'Job Knowledge'
    row_cells[1].text =  'Excellent' if kwargs['job_Knowledge'] == 'Excellent' else 'N/A'
    row_cells[2].text = 'Good' if kwargs['job_Knowledge'] == 'Good' else 'N/A'
    row_cells[3].text = 'Fair' if kwargs['job_Knowledge'] == 'Fair' else 'N/A'
    row_cells[4].text = 'Poor' if kwargs['job_Knowledge'] == 'Poor' else 'N/A'
    row_cells[5].text = kwargs['job_Knowledge_comments']

    row_cells = table.add_row().cells
    row_cells[0].text = 'Productivity'
    row_cells[1].text = 'Excellent' if kwargs['productivity'] == 'Excellent' else 'N/A'
    row_cells[2].text = 'Good' if kwargs['productivity'] == 'Good' else 'N/A'
    row_cells[3].text = 'Fair' if kwargs['productivity'] == 'Fair' else 'N/A'
    row_cells[4].text = 'Poor' if kwargs['productivity'] == 'Poor' else 'N/A'
    row_cells[5].text = kwargs['productivity_comments']

    row_cells = table.add_row().cells
    row_cells[0].text = 'Work Quality'
    row_cells[1].text = 'Excellent' if kwargs['work_quality'] == 'Excellent' else 'N/A'
    row_cells[2].text = 'Good' if kwargs['work_quality'] == 'Good' else 'N/A'
    row_cells[3].text = 'Fair' if kwargs['work_quality'] == 'Fair' else 'N/A'
    row_cells[4].text = 'Poor' if kwargs['work_quality'] == 'Poor' else 'N/A'
    row_cells[5].text = kwargs['work_quality_comments']

    row_cells = table.add_row().cells
    row_cells[0].text = 'Technical Skills'
    row_cells[1].text = 'Excellent' if kwargs['technical_skills'] == 'Excellent' else 'N/A'
    row_cells[2].text = 'Good' if kwargs['technical_skills'] == 'Good' else 'N/A'
    row_cells[3].text = 'Fair' if kwargs['technical_skills'] == 'Fair' else 'N/A'
    row_cells[4].text = 'Poor' if kwargs['technical_skills'] == 'Poor' else 'N/A'
    row_cells[5].text = kwargs['technical_skills_comments']

    row_cells = table.add_row().cells
    row_cells[0].text = 'Work Consistency'
    row_cells[1].text = 'Excellent' if kwargs['work_consistency'] == 'Excellent' else 'N/A'
    row_cells[2].text = 'Good' if kwargs['work_consistency'] == 'Good' else 'N/A'
    row_cells[3].text = 'Fair' if kwargs['work_consistency'] == 'Fair' else 'N/A'
    row_cells[4].text = 'Poor' if kwargs['work_consistency'] == 'Poor' else 'N/A'
    row_cells[5].text = kwargs['work_consistency_comments']

    row_cells = table.add_row().cells
    row_cells[0].text = 'Enthusiasm'
    row_cells[1].text = 'Excellent' if kwargs['enthusiasm'] == 'Excellent' else 'N/A'
    row_cells[2].text = 'Good' if kwargs['enthusiasm'] == 'Good' else 'N/A'
    row_cells[3].text = 'Fair' if kwargs['enthusiasm'] == 'Fair' else 'N/A'
    row_cells[4].text = 'Poor' if kwargs['enthusiasm'] == 'Poor' else 'N/A'
    row_cells[5].text = kwargs['enthusiasm_comments']

    row_cells = table.add_row().cells
    row_cells[0].text = 'Cooperation'
    row_cells[1].text = 'Excellent' if kwargs['cooperation'] == 'Excellent' else 'N/A'
    row_cells[2].text = 'Good' if kwargs['cooperation'] == 'Good' else 'N/A'
    row_cells[3].text = 'Fair' if kwargs['cooperation'] == 'Fair' else 'N/A'
    row_cells[4].text = 'Poor' if kwargs['cooperation'] == 'Poor' else 'N/A'
    row_cells[5].text = kwargs['cooperation_comments']

    row_cells = table.add_row().cells
    row_cells[0].text = 'Attitude/ Behavior'
    row_cells[1].text = 'Excellent' if kwargs['attitude'] == 'Excellent' else 'N/A'
    row_cells[2].text = 'Good' if kwargs['attitude'] == 'Good' else 'N/A'
    row_cells[3].text = 'Fair' if kwargs['attitude'] == 'Fair' else 'N/A'
    row_cells[4].text = 'Poor' if kwargs['attitude'] == 'Poor' else 'N/A'
    row_cells[5].text = kwargs['attitude_comments']

    row_cells = table.add_row().cells
    row_cells[0].text = 'Initiative'
    row_cells[1].text = 'Excellent' if kwargs['initiative'] == 'Excellent' else 'N/A'
    row_cells[2].text = 'Good' if kwargs['initiative'] == 'Good' else 'N/A'
    row_cells[3].text = 'Fair' if kwargs['initiative'] == 'Fair' else 'N/A'
    row_cells[4].text = 'Poor' if kwargs['initiative'] == 'Poor' else 'N/A'
    row_cells[5].text = kwargs['initiative_comments']

    row_cells = table.add_row().cells
    row_cells[0].text = 'Work Relations'
    row_cells[1].text = 'Excellent' if kwargs['work_relations'] == 'Excellent' else 'N/A'
    row_cells[2].text = 'Good' if kwargs['work_relations'] == 'Good' else 'N/A'
    row_cells[3].text = 'Fair' if kwargs['work_relations'] == 'Fair' else 'N/A'
    row_cells[4].text = 'Poor' if kwargs['work_relations'] == 'Poor' else 'N/A'
    row_cells[5].text = kwargs['work_relations_comments']

    row_cells = table.add_row().cells
    row_cells[0].text = 'Creativity'
    row_cells[1].text = 'Excellent' if kwargs['creativity'] == 'Excellent' else 'N/A'
    row_cells[2].text = 'Good' if kwargs['creativity'] == 'Good' else 'N/A'
    row_cells[3].text = 'Fair' if kwargs['creativity'] == 'Fair' else 'N/A'
    row_cells[4].text = 'Poor' if kwargs['creativity'] == 'Poor' else 'N/A'
    row_cells[5].text = kwargs['creativity_comments']

    row_cells = table.add_row().cells
    row_cells[0].text = 'Punctuality'
    row_cells[1].text = 'Excellent' if kwargs['punctuality'] == 'Excellent' else 'N/A'
    row_cells[2].text = 'Good' if kwargs['punctuality'] == 'Good' else 'N/A'
    row_cells[3].text = 'Fair' if kwargs['punctuality'] == 'Fair' else 'N/A'
    row_cells[4].text = 'Poor' if kwargs['punctuality'] == 'Poor' else 'N/A'
    row_cells[5].text = kwargs['punctuality_comments']

    row_cells = table.add_row().cells
    row_cells[0].text = 'Attendance'
    row_cells[1].text = 'Excellent' if kwargs['attendance'] == 'Excellent' else 'N/A'
    row_cells[2].text = 'Good' if kwargs['attendance'] == 'Good' else 'N/A'
    row_cells[3].text = 'Fair' if kwargs['attendance'] == 'Fair' else 'N/A'
    row_cells[4].text = 'Poor' if kwargs['attendance'] == 'Poor' else 'N/A'
    row_cells[5].text = kwargs['attendance_comments']

    row_cells = table.add_row().cells
    row_cells[0].text = 'Dependability'
    row_cells[1].text = 'Excellent' if kwargs['dependability'] == 'Excellent' else 'N/A'
    row_cells[2].text = 'Good' if kwargs['dependability'] == 'Good' else 'N/A'
    row_cells[3].text = 'Fair' if kwargs['dependability'] == 'Fair' else 'N/A'
    row_cells[4].text = 'Poor' if kwargs['dependability'] == 'Poor' else 'N/A'
    row_cells[5].text = kwargs['dependability_comments']

    row_cells = table.add_row().cells
    row_cells[0].text = 'Communication Skills'
    row_cells[1].text = 'Excellent' if kwargs['communication_skills'] == 'Excellent' else 'N/A'
    row_cells[2].text = 'Good' if kwargs['communication_skills'] == 'Good' else 'N/A'
    row_cells[3].text = 'Fair' if kwargs['communication_skills'] == 'Fair' else 'N/A'
    row_cells[4].text = 'Poor' if kwargs['communication_skills'] == 'Poor' else 'N/A'
    row_cells[5].text = kwargs['communication_skills_comments']

    row_cells = table.add_row().cells
    row_cells[0].text = 'Overall Rating'
    row_cells[1].text = 'Excellent' if kwargs['overall_rating'] == 'Excellent' else 'N/A'
    row_cells[2].text = 'Good' if kwargs['overall_rating'] == 'Good' else 'N/A'
    row_cells[3].text = 'Fair' if kwargs['overall_rating'] == 'Fair' else 'N/A'
    row_cells[4].text = 'Poor' if kwargs['overall_rating'] == 'Poor' else 'N/A'
    row_cells[5].text = kwargs['overall_rating_comments']

    paragraph = document.add_paragraph("""Opportunities for Development""")
    paragraph_format = paragraph.paragraph_format
    paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph_format.line_spacing = Pt(18)

    paragraph = document.add_paragraph("""%s""" % kwargs['opportunities'])
    paragraph_format = paragraph.paragraph_format
    paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph_format.line_spacing = Pt(14)

    paragraph = document.add_paragraph("""Reviewers Comments""")
    paragraph_format = paragraph.paragraph_format
    paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph_format.line_spacing = Pt(18)

    paragraph = document.add_paragraph("""%s""" % kwargs['reviewers_comments'])
    paragraph_format = paragraph.paragraph_format
    paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph_format.line_spacing = Pt(14)

    paragraph = document.add_paragraph(""" By signing this form, you confirm that you have discussed this review in detail with your supervisor. Signing this form does not necessarily indicate that you agree with this performance evaluation. """ )
    paragraph_format = paragraph.paragraph_format
    paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph_format.line_spacing = Pt(10)

    document.add_paragraph("""Date : %s""" % kwargs['date2'])
    document.add_paragraph("""Employee Signature : """)
    document.add_picture(kwargs['signature1'], height=Inches(2.0))

    document.add_paragraph("""Date : %s""" % kwargs['date1'])
    document.add_paragraph("""Reviewer Signature : """)
    document.add_picture(kwargs['signature'], height=Inches(2.0))




    file_name_final = '%s_%s.docx' % (kwargs['emp_name'],kwargs['emp_code'])
    document.save('static/docs/' + file_name_final)
    return file_name_final


def send_document_as_mail(**kwargs):
    subject = 'Probation Form - {}'.format(kwargs['emp_name'])

    msg = Message(subject, sender='aigbusiness@aigbusiness.com', recipients=[
        'pkaur@aigbusiness.com'
    ])

    msg.html = """Please find the attached form."""
    with current_app.open_resource('static/docs/' + kwargs['file_name']) as fp:
        msg.attach(
            filename=kwargs['file_name'],
            data=fp.read(),
            content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
    mail.send(msg)

